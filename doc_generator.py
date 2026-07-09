import os
import uuid
from docx import Document
from state import AgentState

OUTPUT_DIR = "./generated_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def document_generator_node(state: AgentState) -> dict:
    doc = Document()
    
    # Add an executive layout style title
    doc.add_heading("Generated Enterprise Artifact", level=0)
    
    # Add assumptions to the generated document for maximum transparency
    if state["plan"].assumptions:
        doc.add_heading("Document Assertions & Scoping Assumptions", level=2)
        for assumption in state["plan"].assumptions:
            doc.add_paragraph(f"• {assumption}", style='List Bullet')
            
    # Compile text sections chronologically
    for section in state["compiled_sections"]:
        doc.add_heading(section["heading"], level=1)
        for para in section["paragraphs"]:
            doc.add_paragraph(para)
            
    filename = f"document_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return {
        "final_docx_path": filepath,
        "logs": [f"Compiled final Word Document saved at {filepath}"]
    }